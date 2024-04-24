import numpy as np
import pandas as pd
import math
from docx import Document
import docx
import os
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Inches,Cm,Pt

def is_nan(num):
    return pd.isna(num)
    # math.isnan(num)
    # return np.isnan(num)


"""Document 对象表示整个文档；Paragrapha 对象标识段落（在输入文档，每一次回车产生新段落）；Run 对象标识相同样式的文本延续"""

"""删除段落"""
def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    # p._p = p._element = None
    paragraph._p = paragraph._element = None


"""合并word中的表格单元格,site1和site2是元组类型"""
def merdge_cell(table,site1,site2):
    a = table.cell(site1[0],site1[1])
    b = table.cell(site2[0],site2[1])
    a.merge(b)


"""在一个文件中，替换两个表格中的内容，从一个表格把某个单元格的内容复制到另外一个表格的某一单元格"""
# rep_site_pair 是需要替换表格的点位，二维元祖，前面是需替换的点位，后面是替换内容所在点位
# rep_table_pair 需要替换的表格位置 前面是需替换的表，后面是替换内容所在表
def replace_table_onesite(document,rep_table_pair,rep_site_pair):
    tables = document.tables
    # 获取点位坐标值
    print('rep_site_pair',rep_site_pair)
    to_row = rep_site_pair[0][0]
    to_line = rep_site_pair[0][1]
    row_data = rep_site_pair[1][0]
    line_data = rep_site_pair[1][1]
    print(to_row,to_line,row_data,line_data)

    # 替换值的cell，因为cell里面可能是多行，一次性全部取出，只要内容，不要格式
    rep_data_cell = tables[rep_table_pair[1]].rows[row_data].cells[line_data].text

    print('rep_data_cell',rep_data_cell)
    # tables[0].rows[to_row].cells[to_line].text = rep_data_cell.text 这种方法会改变原cell里面的格式

    """这里的替换如果不保留格式可以直接替换cell.text,如果要保留格式，必须将cell当作段落来处理，到run级别"""
    cell = tables[rep_table_pair[0]].rows[to_row].cells[to_line]

    # 这里默认只有一个段落，处理run，如果超过1行，其他的行置为空字符串
    runs = cell.paragraphs[0].runs
    for i,run in enumerate(runs):
        if i == 0:
            run.text = rep_data_cell
        if i >= 1:
            run.text = ''

    # paragraph_count = len(cell.paragraphs)
    # run_count = len(cell.paragraphs[0].runs)
    #
    # print('paragraph_count',paragraph_count)
    # print('run_count',run_count)
    #
    # print('old',cell.paragraphs[0].runs[0].text)
    # """必须保证这个cell里面只有一个paragraph，且只有一个run才处理，否则将此文档标记，不做替换"""
    # if paragraph_count == 1 and run_count == 1:
    #     cell.paragraphs[0].runs[0].text = rep_data_cell
    # else:
    #     return -1
#
# for table in document.tables:
#     for row in table.rows:
#         for cell in row.cells:
#             #遍历表格段落内容，回到上个步骤，将cell当作paragraph处理
#             for paragraph in cell.paragraphs:
#                 for run in paragraph.runs:
#                     print(run.text)

"""替换多个，参数以list的形式传入"""
def replace_table_allsite(path_dir,rep_site_pairs):
    files= os.listdir(path_dir)
    fail_list = {}
    # 遍历文件
    for file in files:
        # 判断是否是文件夹，不是文件夹才打开
        if (not os.path.isdir(file)) and (not file.startswith('.')):
            # 打开文件
            file_name = path_dir+"/"+file
            print("file_name:",file_name)
            document = Document(file_name)
            for rep_site_pair in rep_site_pairs:
                res = replace_table_onesite(document,(0,2),rep_site_pair)
                print('res:',res)
                print('--------------')
                print()
            document.save(file_name)

"""将表格中的值用新的文本替换"""
def replace_table_text(table,row_num,cell_num,new_text):
    runs = table.rows[row_num].cells[cell_num].paragraphs[0].runs
    for i,run in enumerate(runs):
        if i == 0:
            run.text = new_text
        if i >= 1:
            run.text = ''

"""在表格中插入文本,英文"""
def insert_table_text(table,row_num,cell_num,new_text):
    table.rows[row_num].cells[cell_num].add_paragraph
    paragraph1 = table.rows[row_num].cells[cell_num].paragraphs[0]
    run = paragraph1.add_run(new_text)
    run.font.size = docx.shared.Pt(12)
    run.font.name = '宋体'

"""在表格中插入文本,中文"""
def insert_table_text_chinese(table,row_num,cell_num,new_text):
    table.rows[row_num].cells[cell_num].add_paragraph
    paragraph1 = table.rows[row_num].cells[cell_num].paragraphs[0]
    run = paragraph1.add_run(new_text)
    run.font.size = docx.shared.Pt(12)
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')


"""替换word段落中的文本"""
def replace_paragraph_text(paragraph,new_text):
    runs = paragraph.runs
    for i,run in enumerate(runs):
        if i == 0:
            run.text = new_text
        if i >= 1:
            run.text = ''

"""设置文档的字体格式"""
def set_style(document):
    # 直接设置中文字体，对中文无效
    paragraph1 = document.add_paragraph()
    run = paragraph1.add_run('aBCDefg这是中文')
    font = run.font
    font.name = '宋体'

    # 方法1 直接修改一个已有样式的所有文字的样式
    style = document.styles['Normal']
    style.font.name = 'Times New Roman' # 必须先设置font.name
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    paragraph1 = document.add_paragraph()
    run = paragraph1.add_run('修改Normal，修改所有字体')

    # 方法2 直接设置文字
    paragraph1 = document.add_paragraph()
    run = paragraph1.add_run('这是直接设置文字的格式 仿宋_GB2312')
    run.font.name ='宋体' # 必须先设置font.name, 只对英文文有效
    # 对中文设置的方法：
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')

    # 方法3 新建样式
    style_song = document.styles.add_style('Song', WD_STYLE_TYPE.CHARACTER)
    style_song.font.name = 'Times New Roman'
    # 设置中文字体
    style_song.element.rPr.rFonts.set(qn('w:eastAsia'), u'黑体')

    paragraph1 = document.add_paragraph() # 添加段落
    paragraph1.add_run(u'方法3 中文字体黑体，英文字体Times New Roman', style='song')


"""从excel获取某几列数据，其中一列作为dic的key，另外的作为value以元组形式输出，参数是excel文件路径，sheet名称"""
def read_excel_from_file(file_path,sheet_name):
    """获得一个pandas.core.frame.DataFrame对象"""
    df = pd.read_excel(io=file_path,sheet_name=sheet_name)
    row_list = df.values.tolist()
    dic = {}
    for row in row_list:
        key = row[0]
        value = (row[3],row[4],row[5],row[7],row[8],row[9],row[11])
        dic[key] = value

    return dic

"""在文档的指定表格，指定位置，插入图片，图片可以是多个"""
def add_picture(table,cell_site,picture_paths,size):
    # 找到要添加的单元格
    cell = table.cell(cell_site[0],cell_site[1])
    for picture_path in picture_paths:
        paragraph = cell.add_paragraph()
        run = paragraph.add_run()
        picture =run.add_picture(picture_path)
        picture.height=Cm(size[0])
        picture.width=Cm(size[1])
        # picture.height=Cm(6.45)
        # picture.width=Cm(7.91)

"""在文档的指定表格，指定位置，替换图片，图片可以是多个"""
def replace_picture(table,cell_site,new_picture_paths,size):
    # 找到要添加的单元格
    cell = table.cell(cell_site[0],cell_site[1])
    delete_paragraph(cell.paragraphs[1])
    add_picture(table,cell_site,new_picture_paths,size)
    # for picture_path in new_picture_paths:
    #     paragraph = cell.add_paragraph()
    #     run = paragraph.add_run()
    #     picture =run.add_picture(picture_path)
    #     picture.height=Cm(size[0])
    #     picture.width=Cm(size[1])
    #     # picture.height=Cm(6.45)
    #     # picture.width=Cm(7.91)


"""path是文件夹的目录 如/Python34/news"""
def read_dir(path):
    # 得到文件夹下的所有文件名称
    files = os.listdir(path)
    s = []

    # 遍历文件
    for file in files:
        # 判断是否是文件夹，不是文件夹才打开
        if not os.path.isdir(file):
            # 打开文件
            f = open(path+"/"+file)
            # 创建迭代器
            iter_f = iter(f)
            str = ""
            # 遍历文件，一行行遍历，读取文本
            for line in iter_f:
                str = str + line
            # 每个文件的文本存到list中
            s.append(str)



if __name__ == '__main__':
    res = is_nan(float(34567789900))
    print(res)
    res_map = replace_table_allsite("/Users/mac/Desktop/测试/400W高的副本",[((0,1),(0,1)),((0,6),(0,6)),((1,1),(1,1)),((1,6),(1,6)),((2,1),(2,1)),((2,5),(2,5))])
    print(res_map)