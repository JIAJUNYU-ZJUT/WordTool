from docx import Document
from docx.oxml.ns import qn
import os
from Utils import read_excel_from_file,insert_table_text,replace_table_text,add_picture,insert_table_text_chinese,is_nan,merdge_cell,delete_paragraph


def single_dir_process(root_path,excel_file_path,sheet_name,save_path,pic_path):
    dic = read_excel_from_file(excel_file_path,sheet_name)
    root_files= os.listdir(root_path)

    print(root_path + "文件夹大小（包含隐藏文件）：",len(root_files))
    t = 0
    # 匹配最后结果，包括匹配的内容
    result = {}
    # 没匹配上的文件列表
    not_pipei_excel = []
    # word段落文本替换有问题
    not_replace = []
    # 图片匹配错误，国标码1为空
    picture_error = []
    # 图片数量缺少或增多
    picture_count_error = []
    for file in root_files:
        result_value = []
        if not file.startswith('.'):
            t = t + 1
            # ------------------第一部分，替换名称----------------------
            ori_file_name = root_path + "/" + file
            document = Document(ori_file_name)
            paragraphs = document.paragraphs

            # 原始点位名称
            ori_dwmc = paragraphs[3].text

            # excel中找到的值 列表
            excel_value = dic.get(ori_dwmc,None)
            if excel_value == None :
                not_pipei_excel.append(file+'')
                # document.save()
                continue
            # 安装类型
            az_type = excel_value[8]
            # 杆件高度
            gj_high = excel_value[9]
            # 挑臂长度
            tb_len = float(excel_value[10])
            # 报警编号
            bjbh = excel_value[0]
            # 国标码1
            gbm1 = excel_value[0]
            # 国标码2
            gbm2 = excel_value[0]
            # 报警编号-点位名称
            bjbh_dwmc = excel_value[3]

            result_value.append('原始点位名称：' + ori_dwmc)
            result_value.append('安装类型：'+az_type)
            result_value.append('杆件高度：' + str(gj_high))
            result_value.append('挑臂长度：' + str(tb_len))
            result_value.append('报警编号：'+bjbh)
            result_value.append('国标码1：'+str(gbm1))
            result_value.append('国标码2：'+str(gbm2))
            result_value.append('报警编号-点位名称：'+bjbh_dwmc)
            # 替换文档中的名称
            if paragraphs[3].runs[0].text == ori_dwmc and paragraphs[23].runs[0].text == ori_dwmc:
                paragraphs[3].runs[0].text = bjbh_dwmc
                paragraphs[23].runs[0].text = bjbh_dwmc
            else:
                not_replace.append(file+'')
                # document.save()
                continue
            # 替换表格中的名称
            table0 = document.tables[0]
            replace_table_text(table0,0,1,bjbh_dwmc)
            # --------------------第一部分-------------------------

            # --------------------第二部分，替换表格内容-------------------------
            # 设备表格
            table1 = document.tables[1]

            if az_type == '借杆':
                gj_xh = '挑杆' + str(tb_len) + '米'
            if az_type == '吊装':
                gj_xh = '吊杆'
            if az_type == '壁挂' and tb_len <= 0.3:
                gj_xh = '原装支架'
            if az_type == '壁挂' and tb_len > 0.3:
                gj_xh = '挑杆' + str(tb_len) + '米'
            if az_type == '立杆':
                gj_xh = '立杆' + str(gj_high) + '米挑杆' + str(tb_len) + '米'
                dyx = 'RVV2*1.5'
                jdx_count = '1'
                replace_table_text(table1,5,4,dyx)
                replace_table_text(table1,6,3,jdx_count)
            replace_table_text(table1,3,4,gj_xh)
            insert_table_text(table1,10,4,'CXBG-1-1-PS-A-C33-YHLZX')
            # 中英文直接设置字体不生效方法不一样
            insert_table_text_chinese(table1,10,5,'海康')
            # ---------------------第二部分------------------------

            # ---------------------第三部分，插入图片-------------------
            # 遍历文件夹下面所有的图片文件,插入图片
            files= os.listdir(pic_path) #得到文件夹下的所有文件名称
            s = []
            f = []

            for pic_file in files: #遍历图片文件
                if pic_file.startswith(bjbh):
                    file_name = pic_path + "/" + pic_file
                    f.append(pic_file + '')
                    s.append(file_name)
            picture_table = document.tables[2]
            # 记录第一个单元格的paragraph数量，后面合并以后只保留这个数量的paragraph，其他的删除
            paragraphs_length = len(picture_table.rows[2].cells[0].paragraphs)
            s_length = len(s)
            # 国标码2为空，合并单元格
            if is_nan(gbm2) and (not is_nan(gbm1)):
                # 合并单元格
                merdge_cell(picture_table,(2,0),(2,1))
                # 删除单元格2中的内容
                paragraphs = picture_table.rows[2].cells[0].paragraphs
                for i,paragraph in enumerate(paragraphs):
                    # 根据上面记录的数量删除多余的paragraph
                    if i >= paragraphs_length:
                        delete_paragraph(paragraph)
                # 图片路径存在，插入图片
                if s_length >= 1:
                    add_picture(picture_table,(2,0),s,(6.45,16.00))
                    if s_length > 1:
                        picture_count_error.append(file+' 图片增多')
                else:
                    picture_count_error.append(file+' 图片缺少')
            elif (not is_nan(gbm2)) and (not is_nan(gbm1)):
                f.sort()
                s.sort()
                if s_length >= 2:
                    add_picture(picture_table,(2,0),[s[0]],(6.45,7.91))
                    add_picture(picture_table,(2,1),[s[1]],(6.45,7.91))
                    if s_length > 2:
                        picture_count_error.append(file+' 图片增多')
                else:
                    if s_length == 1:
                        add_picture(picture_table,(2,0),s,(6.45,7.91))
                    picture_count_error.append(file+' 图片缺少')
            else:
                picture_error.append(file+'')
            result_value.append(f)
            # -----------------第三部分------------------------
            # 保存
            save_file_name = save_path + "/" + file
            document.save(save_file_name)
            result[file+''] = result_value

    print("实际处理的文件个数：", t)
    print("结果dic大小：",len(result))
    print("没有在excel中找到，直接跳过",not_pipei_excel)
    print("没有成功替换点位名称，直接跳过",not_replace)
    print("国标码1为空，文字正常已处理",picture_error)
    print("图片数据异常，文字与图片均处理",picture_count_error)
    for key,value in result.items():
        print(key,value)
    print("**************************")
    print("**************************")
    with open(save_path + "/" + "result.txt","w") as f:
        f.write("文件夹大小（包含隐藏文件）：" + str(len(root_files)) + '\n')
        f.write("实际处理的文件个数：" + str(t) + '\n')
        f.write("结果dic大小：" + str(len(result)) + '\n')
        f.write("没有在excel中找到，直接跳过" + str(not_pipei_excel) + '\n')
        f.write("没有成功替换点位名称，直接跳过" + str(not_replace) + '\n')
        f.write("国标码1为空，文字正常已处理"+ str(picture_error) + '\n')
        f.write("图片数量异常，文字与图片均处理"+ str(picture_count_error) + '\n')
        for key,value in result.items():
            f.write(key + str(value) + '\n')

# 跟single_dir_process相比去掉第二部分操作，图片的操作也有区别
def single_dir_process2(root_path,excel_file_path,sheet_name,save_path,pic_path):
    dic = read_excel_from_file(excel_file_path,sheet_name)
    root_files= os.listdir(root_path)

    print(root_path + "文件夹大小（包含隐藏文件）：",len(root_files))
    t = 0
    # 匹配最后结果，包括匹配的内容
    result = {}
    # 没匹配上的文件列表
    not_pipei_excel = []
    # word段落文本替换有问题
    not_replace = []
    # 图片匹配错误，国标码1为空
    picture_error = []
    # 图片数量缺少或增多
    picture_count_error = []
    for file in root_files:
        result_value = []
        if not file.startswith('.'):
            t = t + 1
            # ------------------第一部分，替换名称----------------------
            ori_file_name = root_path + "/" + file
            document = Document(ori_file_name)
            paragraphs = document.paragraphs

            # 原始点位名称
            ori_dwmc = paragraphs[3].text

            # excel中找到的值 列表
            excel_value = dic.get(ori_dwmc,None)
            if excel_value == None :
                not_pipei_excel.append(file+'')
                # document.save()
                continue
            # 安装类型
            # az_type = excel_value[8]
            # 杆件高度
            # gj_high = excel_value[9]
            # 挑臂长度
            # tb_len = float(excel_value[10])
            # 报警编号
            bjbh = excel_value[0].split('-')[0]
            # 国标码1
            # gbm1 = excel_value[0]
            # 国标码2
            # gbm2 = excel_value[0]
            # 报警编号-点位名称
            bjbh_dwmc = excel_value[0]
            # 设备类型
            sblx = excel_value[1]
            # 计划取电位置
            jhqdwz = excel_value[2]
            result_value.append('设备类型：' + sblx)
            result_value.append('原始点位名称：' + ori_dwmc)
            # result_value.append('安装类型：'+az_type)
            # result_value.append('杆件高度：' + str(gj_high))
            # result_value.append('挑臂长度：' + str(tb_len))
            # result_value.append('报警编号：'+bjbh)
            # result_value.append('国标码1：'+str(gbm1))
            # result_value.append('国标码2：'+str(gbm2))
            # 替换文档中的名称
            if paragraphs[3].runs[0].text == ori_dwmc and paragraphs[23].runs[0].text == ori_dwmc:
                paragraphs[3].runs[0].text = bjbh_dwmc
                paragraphs[23].runs[0].text = bjbh_dwmc
            else:
                not_replace.append(file+'')
                # document.save()
                continue

            table0 = document.tables[0]
            result_value.append('word表格中原点位名称：'+ table0.rows[0].cells[1].paragraphs[0].runs[0].text)
            # 替换表格中的名称
            replace_table_text(table0,0,1,bjbh_dwmc)
            result_value.append('替换的点位名称：' + bjbh_dwmc)

            ori_qdwz = table0.rows[0].cells[6].paragraphs[0].runs[0].text
            result_value.append('word表格中原取电位置：'+ ori_qdwz)
            # 替换取电位置
            replace_table_text(table0,0,6,jhqdwz)
            result_value.append('替换的取电位置：' + jhqdwz)
            result_value.append('取电位置是否相同：' + str(jhqdwz == ori_qdwz))
            if sblx in ('治安监控（800W双目拼接球机）','人脸卡口（800W双摄双云台人脸相机）','人脸卡口（枪球联动摄像机）','高空瞭望（1200W全景AR高空摄像机）'):
                result_value.append('图片应该有：2张')
                tpsl = 2
            elif sblx in ('治安监控（400W低照度球机）','人脸卡口（800W人脸枪机）','高空瞭望（400W高倍率高空摄像机）'):
                result_value.append('图片应该有：1张')
                tpsl = 1
            # --------------------第一部分-------------------------

            # --------------------第二部分，替换表格内容-------------------------
            # 设备表格
            # ---------------------第二部分------------------------

            # ---------------------第三部分，插入图片-------------------
            # 遍历文件夹下面所有的图片文件,插入图片
            files= os.listdir(pic_path) #得到文件夹下的所有文件名称
            s = []
            f = []

            for pic_file in files: #遍历图片文件
                if pic_file.startswith(bjbh):
                    file_name = pic_path + "/" + pic_file
                    f.append(pic_file + '')
                    s.append(file_name)
            picture_table = document.tables[2]
            # 记录第一个单元格的paragraph数量，后面合并以后只保留这个数量的paragraph，其他的删除
            paragraphs_length = len(picture_table.rows[2].cells[0].paragraphs)
            s_length = len(s)
            # 判断设备类型对应的图片数量，合并单元格
            # if is_nan(gbm2) and (not is_nan(gbm1)):
            if sblx in ('治安监控（400W低照度球机）','人脸卡口（800W人脸枪机）','高空瞭望（400W高倍率高空摄像机）'):
                # 合并单元格
                merdge_cell(picture_table,(2,0),(2,1))
                # 删除单元格2中的内容
                paragraphs = picture_table.rows[2].cells[0].paragraphs
                for i,paragraph in enumerate(paragraphs):
                    # 根据上面记录的数量删除多余的paragraph
                    if i >= paragraphs_length:
                        delete_paragraph(paragraph)
                # 图片路径存在，插入图片
                if s_length >= 1:
                    add_picture(picture_table,(2,0),s,(6.46,16.00))
                    if s_length > 1:
                        picture_count_error.append(file+' 图片增多')
                else:
                    picture_count_error.append(file+' 图片缺少')
            # elif (not is_nan(gbm2)) and (not is_nan(gbm1)):
            elif sblx in ('治安监控（800W双目拼接球机）','人脸卡口（800W双摄双云台人脸相机）','人脸卡口（枪球联动摄像机）','高空瞭望（1200W全景AR高空摄像机）'):
                f.sort()
                s.sort()
                if s_length >= 2:
                    add_picture(picture_table,(2,0),[s[0]],(6.46,7.9))
                    add_picture(picture_table,(2,1),[s[1]],(6.46,7.9))
                    if s_length > 2:
                        picture_count_error.append(file+' 图片增多')
                else:
                    if s_length == 1:
                        add_picture(picture_table,(2,0),s,(6.46,7.9))
                    picture_count_error.append(file+' 图片缺少')
            else:
                picture_error.append(file+'')
            result_value.append('图片实际数量：' + str(len(f)) + '张')
            result_value.append('图片数量是否相同：' + str(len(f) == tpsl))
            result_value.append(f)
            # -----------------第三部分------------------------
            # 保存
            save_file_name = save_path + "/" + file
            document.save(save_file_name)
            result[file+''] = result_value

    print("实际处理的文件个数：", t)
    print("结果dic大小：",len(result))
    print("没有在excel中找到，直接跳过",not_pipei_excel)
    print("没有成功替换点位名称，直接跳过",not_replace)
    print("国标码1为空，文字正常已处理",picture_error)
    print("图片数据异常，文字与图片均处理",picture_count_error)
    for key,value in result.items():
        print(key,value)
    print("**************************")
    print("**************************")
    with open(save_path + "/" + "result.txt",mode="w",encoding='utf8') as f:
        f.write("文件夹大小（包含隐藏文件）：" + str(len(root_files)) + '\n')
        f.write("实际处理的文件个数：" + str(t) + '\n')
        f.write("结果dic大小：" + str(len(result)) + '\n')
        f.write("没有在excel中找到，直接跳过" + str(not_pipei_excel) + '\n')
        f.write("没有成功替换点位名称，直接跳过" + str(not_replace) + '\n')
        f.write("国标码1为空，文字正常已处理"+ str(picture_error) + '\n')
        f.write("图片数量异常，文字与图片均处理"+ str(picture_count_error) + '\n')
        for key,value in result.items():
            f.write(key + str(value) + '\n')

"""处理一个目标文件夹下面的所有目录"""
if __name__ == '__main__':
    # dir_path是要处理的所有文件夹的父目录
    dir_path = r"/Users/mac/Desktop/资料待做/仁和" #文件夹目录
    root_paths = os.listdir(dir_path)
    for root_path in root_paths:
        absolute_root_path = dir_path + "/" + root_path
        if os.path.isdir(absolute_root_path) and (not root_path.startswith('.')):
            excel_file_path = r"/Users/mac/Desktop/勘点、单点资料对照（取电）.xlsx"
            sheet_name = "仁和所"
            save_path = r"/Users/mac/Desktop/测试/" + root_path
            if not os.path.exists(save_path):
                os.mkdir(save_path)
            pic_path = r"/Users/mac/Desktop/压缩图片"
            single_dir_process2(absolute_root_path,excel_file_path,sheet_name,absolute_root_path,pic_path)
