from docx import Document
import re

ori_file_name = "/Users/mac/Desktop/新城修改（790）/4（577-726）新城800W人脸/（652-726）新城800W人脸合并-下册.docx"
document = Document(ori_file_name)
paragraphs = document.paragraphs
j = 652
for i,paragrap in enumerate(paragraphs):
    if re.match("文档编号.*",paragrap.text):
        s = paragrap.runs[1].text
        s_exact = s[0:len(s)-5]
        # print(s_exact)
        if j < 10:
            add_str = '0000' + str(j)
        elif j < 100:
            add_str = '000' + str(j)
        else:
            add_str = '00' + str(j)
        paragrap.runs[1].text = s_exact + add_str
        j = j + 1

print(j-1)
document.save(ori_file_name)

document = Document(ori_file_name)
paragraphs = document.paragraphs
k = 0
for i,paragrap in enumerate(paragraphs):
    if re.match("文档编号.*",paragrap.text):
        k = k + 1
        print(paragrap.text)

print(k)