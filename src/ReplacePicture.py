from docx import Document
from docx.oxml.ns import qn
import os
from Utils import replace_picture,read_excel_from_file,insert_table_text,replace_table_text,add_picture,insert_table_text_chinese,is_nan,merdge_cell,delete_paragraph

# 找出文件中点位对应的table，以字典返回
def find_table_by_site_name(document,ori_file_name):
    paragraphs = document.paragraphs
    tbls = document.tables
    site_table_dic = {}
    for i,p in enumerate(paragraphs):
        if i > 1:
            if p.text == '':
                break;
            site_name = p.text.split('-')[0]
            table_number = int((int(p.text.split('\t')[1]) - 1) / 4 * 3 + 2)
            site_table_dic[site_name] = (tbls[table_number],ori_file_name)
    return site_table_dic

# 根据图片文件找出所有需要修改的点位，以dic返回
def find_site_picture(pic_dir):
    files= os.listdir(pic_dir) #得到文件夹下的所有文件名称
    all_replace_site = {}
    # print(files)
    for pic_file in files: #遍历图片文件
        # print(pic_file)
        if not pic_file.startswith('.DS_Store'):
            site_name = pic_file.split('-')[0]
            if all_replace_site.get(site_name,None):
                l = all_replace_site[site_name]
                l.append(pic_dir + pic_file)
                l.sort()
                all_replace_site[site_name] = l
            else:
                all_replace_site[site_name] = [pic_dir + pic_file]
    return all_replace_site

def replace_picture_by_site_name(site_table,site_picture,dic_replace_result):
    for site_name in site_picture.keys():
        # print(site_name)
        if site_table.get(site_name,None):
            pic_list = site_picture[site_name]
            # 根据图片数量判断要替换的数量
            if len(pic_list) == 2:
                add_picture(site_table[site_name][0],(2,0),[pic_list[0]],(6.44,7.89))
                add_picture(site_table[site_name][0],(2,1),[pic_list[1]],(6.44,7.89))
                dic_replace_result[site_name].append("替换 %s 文件中的 %s 点位2张图片" % (site_table[site_name][1],site_name))
            else:
                add_picture(site_table[site_name][0],(2,0),[pic_list[0]],(6.44,16.00))
                dic_replace_result[site_name].append("替换 %s 文件中的 %s 点位1张图片" % (site_table[site_name][1],site_name))



# 图片文件所在文件夹
pic_dir = '/Users/mac/Desktop/单点资料需修改0428/仁和人脸-压缩/'
# 点位图片对应图
site_picture = find_site_picture(pic_dir)
print("点位图片",site_picture)
# 替换结果表
dic_replace_result = {k : [] for k in site_picture.keys()}
# print(dic_replace_result)
# word文件所在文件夹
file_dir = '/Users/mac/Desktop/单点资料需修改0428/3-2.仁和-可打印/'
#得到文件夹下的所有文件名称
files= os.listdir(file_dir)
# 遍历处理word文件
for word_file in files:
    if not word_file.startswith('.'):
        # print(word_file)
        # 文件全路径名
        ori_file_path = file_dir + word_file
        # 文件名
        ori_file_name = word_file
        # 保存文件全路径名
        save_file_path = file_dir + "ceshi" +word_file
        # 打开文件
        document = Document(ori_file_path)
        # 获取文件点位table对应图
        site_table = find_table_by_site_name(document,ori_file_name)
        print("点位table",site_table)
        # 根据点位table和点位图片结果进行替换替换
        replace_picture_by_site_name(site_table,site_picture,dic_replace_result)
        # 保存文件
        # document.save(save_file_path)
        document.save(ori_file_path)
for k,v in  dic_replace_result.items():
    print(k,v)
