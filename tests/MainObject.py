from docx import Document
import os
from Utils import read_excel_from_file,insert_table_text,replace_table_text,add_picture

dic = read_excel_from_file(r"/Users/mac/Desktop/瓶窑单点资料（299）/删减-项目进度跟踪表3.13.xlsx","瓶窑")

root_path = r"/Users/mac/Desktop/瓶窑单点资料（299）/1200W高空（5）/" #文件夹目录
root_files= os.listdir(root_path)

print("文件夹大小（包含隐藏文件）：",len(root_files))
t = 0
j = 0
# 匹配结果
result = {}
not_pipei = []
for file in root_files:
    result_value = []
    j = j + 1
    if not file.startswith('.'):
        t = t + 1
        # ------------------第一部分，替换名称----------------------
        ori_file_name = root_path + file
        document = Document(ori_file_name)
        value = ''
        az_type = ''
        gj_high = ''
        tb_len = 0
        bjbh = ''
        rr = None
        for i,para in enumerate(document.paragraphs):
            if i == 3:
                key = para.text
                rr = dic.get(key,None)
                if rr:
                    bjbh = rr[4]
                    dw_name = rr[3]
                    az_type = rr[0]
                    gj_high = rr[1]
                    tb_len = float(rr[2])
                    result_value.append('原始点位名称：'+key)
                    result_value.append('报警编号：'+bjbh)
                    result_value.append('报警编号-点位名称：'+dw_name)
                    result_value.append('安装类型：'+az_type)
                    result_value.append('杆件高度：'+ str(gj_high))
                    result_value.append('挑臂长度：'+ str(tb_len))
                else:
                    not_pipei.append(file+'')
            if (i == 3 or i == 23) and rr:
                para.runs[0].text = dw_name
        table = document.tables[0]
        if rr:
            table.rows[0].cells[1].paragraphs[0].runs[0].text = dw_name
        # --------------------第一部分-------------------------

        # --------------------第二部分，替换表格内容-------------------------
        # 设备表格
            sb_table = document.tables[1]

            if az_type == '借杆':
                gj_xh = '挑杆' + str(tb_len) + '米'
            if az_type == '吊装':
                gj_xh = '吊杆'
            if az_type == '壁挂' and tb_len <= 0.3:
                gj_xh = '原装支架'
            if az_type == '壁挂' and tb_len > 0.3:
                gj_xh = '挑杆' + str(tb_len) + '米'
            if az_type == '立杆':
                gj_xh = '立杆' + str(gj_high) + '米挑杆' + str(tb_len) + '米'
                dyx = 'RVV2*1.5'
                jdx_count = '1'
                replace_table_text(sb_table,5,4,dyx)
                replace_table_text(sb_table,6,3,jdx_count)
            replace_table_text(sb_table,3,4,gj_xh)
            insert_table_text(sb_table,10,4,'CXBG-1-1-PS-A-C33-YHLZX')
            insert_table_text(sb_table,10,5,'海康')
        # ---------------------第二部分------------------------

        # ---------------------第三部分，插入图片-------------------
        # 遍历文件夹下面所有的图片文件,插入图片
            path = "/Users/mac/Desktop/瓶窑单点资料（299）/瓶窑系统截图" #文件夹目录
            files= os.listdir(path) #得到文件夹下的所有文件名称
            s = []
            f = []

            for pic_file in files: #遍历文件夹
                if pic_file.startswith(bjbh):
                    file_name = "/Users/mac/Desktop/瓶窑单点资料（299）/瓶窑系统截图/" + pic_file
                    f.append(pic_file + '')
                    s.append(file_name)
            picture_table = document.tables[2]
            add_picture(picture_table,(2,0),s,(6.45,7.91))
            result_value.append(f)
        # -----------------第三部分------------------------
        # 保存
        document.save(ori_file_name)
        result[file+''] = result_value

print("实际处理的文件个数：",t)
print("结果dic大小：",len(result))
print(not_pipei)
for key,value in result.items():
    print(key,value)