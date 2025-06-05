# from docx import Document
import docx
# 因为com.souche.word不是source_root，idea不识别，只会识别sourceroot python文件夹下面的资源，所以会有下划线，虽然运行不报错（前提是在同一个文件夹下面，是运行前idea帮助添加了sys.path.append）
# content root是项目根,这个目录idea也会自动加入sys.path，跟source root 一样，这个在preferences -console - python console中设置
# python.com.souche.word.Utils这种全路径的方式就不会错，因为python是sourceroot，根据这个路径能找到目标
# mark sourceroot这种方式在idea运行是可以的，但是如果在控制台运行，就没有工具在运行前加到 sys.path 里了，所以还是在代码中要加上 sys.path.append(XXX),手动添加
# from python.com.souche.word.Utils import is_nan 这一行如果不报错就要加 sys.path.append(/Users/mac/IdeaProjects/pythonTest/python)
# Utils import is_nan 这一行如果不报错就要加 sys.path.append(/Users/mac/IdeaProjects/pythonTest/python/com/souche/word)
# 加一个标记
from Utils import is_nan,read_excel_from_file
import sys,os

from docx.shared import Inches,Cm,Pt

# print(is_nan(100))

# to_rep_site 是需要替换表格的点位，二位元祖，rep_site是替换内容所在点位

def find_site_picture(pic_dir):
    files= os.listdir(pic_dir) #得到文件夹下的所有文件名称
    all_replace_site = {}
    # print(files)
    for pic_file in files: #遍历图片文件
        # print(pic_file)
        if not pic_file.startswith('.DS_Store'):
            site_name = pic_file.split('-')[0]
            if all_replace_site.get(site_name,None):
                l = all_replace_site[site_name]
                l.append(pic_dir + pic_file)
                l.sort()
                all_replace_site[site_name] = l
            else:
                all_replace_site[site_name] = [pic_dir + pic_file]
    return all_replace_site


def replace_table(file_name,to_rep_site,rep_site):
    document=docx.Document(file_name)
    tables = document.tables

    # 获取点位坐标值
    to_row = to_rep_site[0]
    to_line = to_rep_site[1]

    row_data = rep_site[0]
    line_data = rep_site[1]

    # 值的cell，因为cell里面可能是多行
    rep_data_cell = tables[0].rows[row_data].cells[line_data]
    rep_data_run = tables[0].rows[row_data].cells[line_data].paragraphs[0].runs[0]

    print(rep_data_cell.text)
    print('---------------')
    print(rep_data_run.text)

    # tables[0].rows[to_row].cells[to_line].text = rep_data_cell.text
    # document.save(file_name)


# replace_table("/Users/mac/Desktop/测试/400W高的副本/安溪东王庙北制高点.docx",(0,1),(1,0))
# document=Document("/Users/mac/Desktop/测试/400W高的副本/安溪东王庙北制高点.docx")
# table = document.tables[2]
# row = table.rows[2]
# for cell in row.cells:
#     print(cell.text)
    # print(cell)

# document = Document("/Users/mac/Desktop/瓶窑单点资料（299）/400W人脸（98）的副本/2023NYHQTWGC（GABJ）-DDYSD-PYPCS-00080.docx")
# path = r"/Users/mac/Desktop/瓶窑单点资料（299）/瓶窑系统截图/" #文件夹目录
# files= os.listdir(path) #得到文件夹下的所有文件名称
# s = []
#
# for file in files: #遍历文件夹
#     if file.startswith("瓶1330119"):
#         file_name = "/Users/mac/Desktop/瓶窑单点资料（299）/瓶窑系统截图/" + file
#         s.append(file_name)
# print(s)
#
# InsertPicture.add_picture(document,2,(2,0),s)
# document.save("/Users/mac/Desktop/瓶窑单点资料（299）/400W人脸（98）的副本/2023NYHQTWGC（GABJ）-DDYSD-PYPCS-00080.docx")


# document = Document("/Users/mac/IdeaProjects/pythonTest/python/com/souche/word/PYPCS-00001.docx")
# # table = document.tables[1]
# for i,paragraph in enumerate(document.paragraphs):
#     print(str(i) + paragraph.text)

#
# print(document.paragraphs[3].text)
# print(document.paragraphs[23].text)

# ReplaceTest.insert_table_text(table,10,4,'CXBG-1-1-PS-A-C33-YHLZX')
# ReplaceTest.insert_table_text_chinese(table,10,5,'海康')

# paragraphs_length = len(table.rows[2].cells[0].paragraphs)
#
# # 合并单元格
# ReplaceTest.merdge_cell(table,(2,0),(2,1))
#
# # 删除单元格2中的内容
# paragraphs = table.rows[2].cells[0].paragraphs
# for i,paragraph in enumerate(paragraphs):
#     if i >= paragraphs_length:
#         ReplaceTest.delete_paragraph(paragraph)
#
# # 添加图片
# InsertPicture.add_picture(table,(2,0),["/Users/mac/IdeaProjects/pythonTest/python/com/souche/word/pictest.jpg"],(6.45,16.00))

# document.save("/Users/mac/IdeaProjects/pythonTest/python/com/souche/word/PYPCS-00001-a.docx")

# l1 = ['瓶1330137-双红桥街与下金路西B_2024_ 3_15_ 9_35_ 1_31.jpg', '瓶1330137-双红桥街与下金路西A_2024_ 3_15_ 9_34_57_111.jpg']
# print(l1)
# l1.sort()
# print(l1)

# dir_path = r"/Users/mac/Desktop/tt.docx"
# document = Document(dir_path)
# replace_picture(document.tables[0],(2,0),["/Users/mac/Desktop/黄.jpg"],(6.45,16.00))
# document.save("/Users/mac/Desktop/tt1.docx")

# root_paths = os.listdir(dir_path)
# for root_path in root_paths:
#     print(root_path)
    # absolute_root_path = dir_path + "/" + root_path
    # if os.path.isdir(absolute_root_path) and (not root_path.startswith('.')):
    #     print(absolute_root_path)
# dic1 = dict({'a':1},{'pay_money':'pay_money','time':'trade_time'})
# dic1 = dict(**{'a':1},**{'a':'pay_money','time':'trade_time'})
# dic2 = dict([('x', 5), ('y', -5)], z=8)
# print(dic1)

# document=Document("/Users/mac/Desktop/12.docx")
# table = document.tables[0]
# row = table.rows[7]
# i = 0
# for cell in row.cells:
#     i= i + 1
#     print(i)
#     print(cell.text)
#     print(cell.width)


# print(cell)
#
# res = read_excel_from_file('/Users/mac/Desktop/仁和/勘点、单点资料对照.xlsx','仁和所')
# print(res)
if __name__ == '__main__':
    document = docx.Document('/Users/mac/Desktop/资料待做/仁和/资料-400W低照度球机/YHQTWGCXMEQ（GABJ）-DDYSD-RHPCS-00008.docx')
    table0 = document.tables[0]
    run = table0.rows[0].cells[6].paragraphs[0].runs[0]
    print(run.text)
    # cells = table0.rows[0].cells
    # for j,cell in enumerate(cells):
    #     print(j)
    #     print(cell.paragraphs[0].runs[0].text)
    # run1 = table0.rows[0].cells[1].paragraphs[0].runs[0]
    # run2 = table0.rows[1].cells[1].paragraphs[0].runs[0]
    # run3 = table0.rows[2].cells[3].paragraphs[0].runs[0]
    # print(run1.text)
    # print(run2.text)
    # print(run3.text)
    # paragraphs = document.paragraphs
    #
    # ori_dwmc = paragraphs[3].text
    # print(ori_dwmc)
    # print(paragraphs[3].runs[0].text)
    # print(paragraphs[23].runs[0].text)
    # document.core_properties
    # secs = document.sections
    # print(len(document.part))
    # document.settings.odd_and_even_pages_header_footer = True
    # document.sections[0].even_page_header.paragraphs[0].text = "这是偶数页页眉"
    # document.sections[0].header.paragraphs[0].text = "这是奇数页页眉"
    #
    # # for j,sec in enumerate(secs):
    # #     print(j)
    # #     sec.header.paragraphs[0].text = "这是第%s页眉"%(str(j))
    # document.save('/Users/mac/Desktop/测试/t.docx')

    # document.core_properties.pages
    # paragraphs = document.paragraphs
    # for i,para in enumerate(paragraphs):
    #     print(str(i) + para.text)
    #     print(len(para.text))
