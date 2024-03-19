from docx import Document
# from docx.enum.section import WD_SECTION
import docx

from docx import Document
from docx.shared import Inches

class Item:
    def __init__(self,qty,id,desc):
        self.qty = qty
        self.id = id
        self.desc = desc

recordset = set()

recordset.add(Item(10,3,'wos'))
recordset.add(Item(20,2,'dd'))
recordset.add(Item(23,1,'iu'))

document = Document()

document.add_heading('Document Title', 0)

p = document.add_paragraph('A plain paragraph having some ')
p.add_run('bold').bold = True #加粗
p.add_run(' and some ')
p.add_run('italic.').italic = True #倾斜
run = p.add_run('我是')
run.font.size = docx.shared.Pt(12)
run.font.name = '宋体'

document.add_heading('Heading, level 1', level=1)
document.add_paragraph('Intense quote', style='IntenseQuote')

document.add_paragraph(
    'first item in unordered list', style='ListBullet'
)
document.add_paragraph(
    'first item in ordered list', style='ListNumber'
)


table = document.add_table(rows=6, cols=6)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Qty'
hdr_cells[1].text = 'Id'
hdr_cells[2].text = 'Desc'
for item in recordset:
    row_cells = table.add_row().cells
    row_cells[0].text = str(item.qty)
    row_cells[1].text = str(item.id)
    row_cells[2].text = item.desc

a = table.cell(2,3)
b = table.cell(4,5)
a.merge(b)
# 在末尾增加一页
# document.add_page_break()
document.sections[-1]
document.add_section()
document.save('一页1.docx')

